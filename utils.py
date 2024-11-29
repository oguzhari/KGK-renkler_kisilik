import streamlit as st
from docx import Document
from datetime import datetime, timedelta
from pydrive.auth import GoogleAuth
from pydrive.drive import GoogleDrive
import random
import string
import color_tone
import send_mail

st.config.set_option("theme.primaryColor", "white")

css_code = """
<style>
    .content-container {
        border-top: 3px solid black;
        border-bottom: 3px solid black;
    }
    .row-divider {
        border-bottom: 1px solid gray;
        margin: 10px 0;
    }
</style>
"""

button_css = """
<style>
    .stButton button {
        background-color: #233b77; /* Green background */
        border: none; /* Remove borders */
        color: white; /* White text */
        padding: 15px 32px; /* Some padding */
        text-align: center; /* Centered text */
        text-decoration: none; /* Remove underline */
        display: inline-block; /* Get the element to line up correctly */
        font-size: 16px; /* Increase font size */
        margin: 4px 2px; /* Some margin */
        cursor: pointer; /* Pointer/hand icon */
        transition-duration: 0.4s; /* 0.4 second transition effect to hover state */
    }

    .stButton button:hover {
        background-color: white; /* White background on hover */
        color: black; /* Black text on hover */
        border: 2px solid #233b77; /* Green border on hover */
    }
</style>
"""

# Apply the CSS styles
st.markdown(button_css, unsafe_allow_html=True)

# CSS kodlarını uygulamaya ekleyin
st.markdown(css_code, unsafe_allow_html=True)

# Metinler
a = "1. Aşağıdakilerden hangisi sizi daha iyi anlatır?"
a1 = "Güçlü, kararlı, girişken ve doğuştan liderim. Kimseye minnet etmem; düşer kalkar ve yoluma devam ederim."
a2 = "Hayata anlamlı renkler katar, eğlenceyi severim. Ömür boyu herkesin mutlu ve neşeli olmasını dilerim"
a3 = "Her anımı huzurlu ve sakin geçirmek isterim. Kavga-gürültü sevmem, işlerimde en kolay yolu seçerim."
a4 = "Her şeyin mükemmel, düzgün ve kusursuz olmasını isterim. İlişkilerimde saygılı ve mesafeli olmayı severim."
b = "2. Genellikle hangi tempoda ve nasıl konuşursunuz?"
b1 = "Hızlı ve sonuca yönelik."
b2 = "Çok hızlı, heyecanlı ve eğlenceli."
b3 = "Daha yavaş ve sakin."
b4 = "Normal ve söyleyeceklerimi aklımda tutarak."
c = "3. Bir işe motive olmanızı sağlayan en önemli unsur hangisidir?"
c1 = "Sonuçları düşünmek."
c2 = "Onaylanmak, takdir edilmek."
c3 = "Gruptaki arkadaşlarımın desteği."
c4 = "Etkinlik, düzen ve disiplin."
d = "4. Çalışma tarzınız hangisine uygundur?"
d1 = "Yoğun ve hızlıyımdır. Aynı anda birkaç işi bir arada yapabilirim."
d2 = "Özgür bir ortamda çalışırım. İnsan ilişkileri odaklıyımdır."
d3 = "Ön planda olmayan ama gruba her zaman destek veren bir yapım vardır."
d4 = "Ayrıntıları önemserim ve tek bir konuya odaklanarak çalışırım."
e = "5. Çalışma temponuzu nasıl değerlendiriyorsunuz?"
e1 = "Hızlı bir tempoda çalışır, çabuk karar almayı severim."
e2 = "İşlerin rutin ve sıkıcı olmadığı ortamlarda yüksek motivasyonla çalışırım."
e3 = "Nadiren aceleciyimdir. Geç de olsa üstlendiğim işi bitiririm."
e4 = "Ayrıntılı düşünerek karar veririm, ağır ama iş bitirici bir tempoyla çalışırım."
f = "6. Hangisi sizi daha çok rahatsız eder?"
f1 = "Zaman israfı ve işlerin gecikmesi."
f2 = "Tekrar gerektiren işler ve monotonluk."
f3 = "Çalışma ortamı ve anlaşmazlıklar."
f4 = "Yanılmak ve yapılan hatanın tekrarlanması."
g = "7. Bulunduğunuz gruplarda hangi konumda daha başarılı olursunuz?"
g1 = "Olaylara yön veren ve otoriteyi kullanan."
g2 = "İnsanları motive eden ve neşelendiren."
g3 = "Uzlaştırıcı ve grup içindeki uyumu sağlayan"
g4 = "Bilgi sağlayıcı, araştırıcı ve olayları takip eden."
h = "8. Hangisi sizi daha çok strese sokar?"
h1 = "Olayların üzerindeki güç ve kontrolümün azaldığını hissetmek."
h2 = "Sıkıcı, rutin işler yapılan bir ortamda bulunmak."
h3 = "Beni aşacağını düşündüğüm sorumluluklar üstlenmek."
h4 = "Düzensiz ortamlar ve eksik yapılan işler."
i = (
    "9. Bir öğrenci olsanız ve öğretmeniniz sınav kağıdınızı ikinci defa incelediğinde puanınızın arttığını söylese, "
    "nası bir tepki verirsiniz?"
)
i1 = "Bunu zaten hakettiğimi düşünürüm."
i2 = "Çok sevinirim ve sevincimi belli ederim."
i3 = "Hocama teşekkür eder ve saygı duyarım."
i4 = "Hocamın nerede hata yaptığını merak eder, kağıdımı görmek isterim."
j = "10. Saatler sürecek bir iş toplantısına katılmanız gerektiğinde aşağıdakilerden hangisini benimsersiniz?"
j1 = (
    "Toplantı başında genellikle konunun ana hatları konuşulduğu için biraz geç girerim. Sonucun belli olmasından "
    "hemen sonra da çıkmayı tercih ederim."
)
j2 = (
    "Toplantı eğlenceli bir şekilde devam ederse sonuna kadar kalırım. Toplantı sıkıcı olmaya "
    "başladığında erken çıkarım."
)
j3 = "Toplantının huzur içinde geçmesi ve güzel kararlar alınması için üstüme düşeni yaparım."
j4 = (
    "Toplantıya tam vaktinde veya vaktinden önce gelirim. Toplantı esnasında notlar alır, "
    "sonunda biraz kalarak değerlendirme yaparım."
)
k = "11. Kendinizde gördüğünüz en zayıf yönünüz hangisidir?"
k1 = "İşler zamanında ve istediğim gibi yapılmadığında sinirlenmek."
k2 = "Düzensiz, dağınık ve plansız olmak."
k3 = "Kimseye 'hayır' diyememek, başkalarının işine koşarken kendi işimi aksatmak."
k4 = "Her şeyin kusursuz, mükemmel olmasını istemek; insanlarda bunu görmediğimde sinirlenmek."
m = "12. Kendinizde gördüğünüz en güçlü yönünüz hangisidir?"
m1 = "Kısa sürede karar alıp hemen harekete geçmem."
m2 = "Girdiğim ortama neşe ve heyecan katabilmem."
m3 = "Her türlü ortama uyum sağlamam ve çatışmaları önleme gayretim."
m4 = "Her şeyi planlı, programlı ve düzenli yapmam."
n = "13. Aşağıdaki ifadelerden hangisi sizi daha iyi tanımlar?"
n1 = "Güçlü, kararlı, otoriter ve yönlendirici."
n2 = "Popüler, neşeli, sevimli ve muzip."
n3 = "Barışçıl, sevecen, uyumlu ve sakin."
n4 = "Tertipli, düzenli, disiplinli ve planlı."
o = "14. Çalışma masanızda nelere dikkat edersiniz?"
o1 = "Öncelikli işlerime göre düzenlenmiş, sade bir masayı tercih ederim."
o2 = "İnsanlara karma karışık gelen ama benim aradığım her şeyi kolayca bulduğum bir masada çalışırım."
o3 = (
    "Önce masamın üzerine gerekli olan her türlü araç gereci koyarım, "
    "çünkü sık sık kalkarak enerjimi harcamak istemem."
)
o4 = "İyi bir iş çıkarmam için masam son derece derli ve düzenli olmalıdır."
p = "15. Ertesi gün çözülmesi gereken bir problem varsa, o akşamki ruh haliniz nasıl olur?"
p1 = "Büyük bir tedirginlik duymam, çünkü ertesi gün olması gerektiği gibi yapacağımdan eminimdir."
p2 = "Çok tedirginlik duymam, çünkü nasıl olsa işler bir şekilde halolacaktır."
p3 = "Sorumluluğun üstümde olmasından dolayı tedirginlik duyarım."
p4 = "Tedirginlik duyarım ve gecenin büyük bir bölümünde problemin nasıl çözüleceğiyle ilgili planlar yaparım."

# Analiz Sonuçları
sari = """
Sarı kişilikli bireyler, enerjik, neşeli ve sosyal olarak tanımlanır. Bu kişilik tipine sahip olanlar, genellikle eğlenceli ve cana yakındır ve insanlarla kolayca iletişim kurarlar. Girişken ve hareketli olmaları, çevrelerindeki insanlar üzerinde pozitif bir etki yaratır. El, kol ve mimiklerini etkili bir şekilde kullanarak, güçlü bir iletişim kurarlar. Enerjilerini insanlardan alırlar ve konuşmayı çok severler.

Sarı kişilikli bireylerin hızlı düşünme ve hazırcevap olma becerileri vardır. İkna kabiliyetleri yüksek olduğu için, fikirlerini başkalarına kolayca aktarabilirler. Sorunlara ilginç ve yaratıcı çözümler bulma konusunda yeteneklidirler ve enerjik tutumlarıyla mutluluk ve neşe yayabilirler. Meraklı ve çocuk ruhlu olmaları, sürekli yeni fikirler üretmelerine ve yeni deneyimler yaşamalarına olanak tanır. İnsanlarla kolayca kaynaşarak, hızlı ve samimi arkadaşlıklar kurabilirler.

Ancak sarı kişilikli bireylerin de negatif yönleri bulunmaktadır. Çalışma alanları dağınık olabilir ve dikkatsiz davranarak, sık sık bir şeyleri döküp saçabilirler. Unutkan olabilirler ve işe odaklanmada zorluk yaşayabilirler. Ayrıntılar üzerine düşünmeyi sevmezler ve genellikle birine ihtiyaç duyarlar. Plan ve program yapmazlar, spontane yaşarlar. Olayları anlatırken abartma eğiliminde olabilirler. İşi başlatır ama sonunu getirmez ve işlerini çoğu zaman ertelerler.

Sarı kişilikli bireylerin başarılı olmaları için, daha düzenli ve planlı olmaya özen göstermeleri gerekmektedir. Ayrıca, işlerini tamamlamak ve ertelememek için önceliklerini belirlemeye ve zaman yönetimi becerilerini geliştirmeye ihtiyaç duyarlar. Enerjilerini ve sosyal becerilerini, ekip çalışmalarında ve insan ilişkilerinde olumlu bir şekilde kullanarak, başarılı ve mutlu olmaları mümkündür.
"""

mavi = """
Mavi kişilik özellikleri, ciddi, ağırbaşlı ve düşünceli bireyleri tanımlar. Bu kişilik tipine sahip olanlar, mükemmeliyetçidir ve detaylara büyük önem verirler. İş hayatında düzenli, planlı ve programlı hareket ederler. Hassas ve derin düşünceli olmalarıyla, riskleri önceden görebilme ve olası sorunlara karşı hazırlıklı olma becerilerine sahiptirler. İdealist olan bu bireyler, kurallara ve prensiplere büyük önem verirler. Araştırmacı ve meraklı bir yapıya sahip oldukları için, sürekli olarak yeni bilgiler edinme ve öğrenme süreçlerinde etkin rol alırlar.

Mavi kişilikli bireylerin bazı negatif yönleri de bulunmaktadır. Detaylara fazla takılma ve karar vermekte zorlanma gibi özellikler bunlardandır. Hata yapmaktan çekinirler ve hayata karamsar bir bakış açısıyla yaklaşabilirler. Aşırı kuralcı olma ve depresyona girebilme eğilimi, yaşamlarını olumsuz yönde etkileyebilir. Olayların olumsuz yönlerini görmeye eğilimli oldukları için, kaygı düzeyleri yüksek olabilir. Mavi kişilikli bireyleri motive etmek için, işin ayrıntılarıyla anlatmak ve planlı, programlı bir şekilde işin yürütüleceğini vurgulamak önemlidir.

Mavi kişilikli bireylerin, özellikle ekip çalışmalarında dikkat etmesi gereken noktalar bulunmaktadır. İşbirliğine daha açık olmaları ve diğer kişilik tipleriyle daha uyumlu çalışabilmeleri için, daha esnek ve adaptasyon yeteneklerini geliştirmeleri gerekmektedir. Ayrıca, detaylarda fazla takılmamaya ve daha hızlı karar alabilme becerisine odaklanmaları önemlidir. Bu sayede, mavi kişilikli bireyler, ekiplerinde daha etkin ve başarılı bir rol üstlenebilirler.
"""

kirmizi = """
Kırmızı kişilik özellikleri, enerjik, dışa dönük ve güçlü karakterlere sahip bireyleri tanımlar. Bu kişilik tipinde olan bireyler, doğal liderler olarak görülür ve takımlarında öncü rol üstlenmeye meyillidirler. İş hayatında sonuç odaklıdırlar ve kararlılıkla hedeflerine yönelirler. Bağımsız hareket etmeye yatkın olan bu kişiler, rekabeti sever ve zaman yönetimi konusunda başarılıdırlar. Hedeflerini belirleyip, gerçekleştirme konusunda istekli ve hırslıdırlar. Sorunlara karşı çözüm odaklı yaklaşımları ve analitik düşünme yetenekleri sayesinde etkili ve başarılı sonuçlar elde edebilirler.

Ancak kırmızı kişilik özelliklerine sahip bireylerin bazı negatif yönleri de bulunmaktadır. Bunlar arasında hep haklı olduklarını düşünme, ani karar alma ve değerlendirme yapamama gibi özellikler bulunur. Eleştirilmekten hoşlanmazlar ve yardım almaya açık olmayabilirler. Dik başlı ve inatçı olmaları nedeniyle, iş birliği ve uyum sağlamada zorluk yaşayabilirler. Kırmızı kişilikli bireyleri motive etmek için, işin sonucu ile ilgili sözler söylemek etkili olabilir. Başarılarını ve sonuç odaklı yaklaşımlarını takdir etmek ve onların liderlik özelliklerini ön plana çıkarmak da önemlidir.

Kırmızı kişilikli bireylerin özellikle ekip çalışmalarında dikkat etmesi gereken noktalar bulunmaktadır. İşbirliğine daha açık olmaları ve diğer kişilik tipleriyle daha uyumlu çalışabilmeleri için, eleştiriye daha açık olmaları ve başkalarının fikirlerine değer verme becerilerini geliştirmeleri gerekmektedir. Bu sayede, kırmızı kişilikli bireyler, ekiplerinde daha etkin ve başarılı liderler olabilirler.
"""

yesil = """
Yeşil kişilik özellikleri, genellikle barışçıl, sakin ve uyumlu bireylerin davranışlarını tanımlar. Bu kişilik tipine sahip olanlar, kendileriyle ve çevreleriyle barışık yaşarlar ve soğukkanlı bir tutum sergilerler. Çatışma ve anlaşmazlıklardan kaçınır, doğal arabulucu rollerini üstlenerek, ortamın dengesini sağlamaya çalışırlar. Sabırlı, uysal ve güvenilir oldukları için, insanlarla iyi ilişkiler kurar ve çevrelerinde sevilen bireyler olurlar. Hislerini iyi saklar ve genellikle duygularını ifade etmekte zorlanırlar.

Yeşil kişilikli bireylerin negatif yönleri de bulunmaktadır. Çekingendirler ve düşük enerjili olabilirler. Sorumluluk almaktan çekinir ve hayır demekte zorlanırlar. Karar verme süreçlerinde zorlanabilir ve duygularını yansıtmama eğilimi gösterebilirler. İşlerini ertelemeye meyillidirler ve değişime zor uyum sağlarlar. Monoton ve düzenli yaşamlarını sürdürmeyi tercih ederler.

Yeşil kişilikli bireylerin çalışma ortamında dikkat etmeleri gereken bazı noktalar bulunmaktadır. Daha fazla sorumluluk alarak ve karar verme süreçlerinde daha etkin olmaları, kişisel ve profesyonel gelişimlerine katkı sağlayacaktır. İnsanlarla iyi ilişkiler kurma becerilerini kullanarak, ekip çalışmalarında daha etkin ve başarılı olabilirler. Duygularını daha açık bir şekilde ifade etmeye çalışarak, daha iyi anlaşılmalarını ve gereksinimlerini daha net ortaya koymalarını sağlar. Bu sayede, yeşil kişilikli bireyler hem kendi yaşamlarında hem de çalışma ortamlarında daha başarılı ve mutlu olabilirler.
"""

kirmizi_sari = """
Kırmızı ve sarı kişilik kombinasyonu hem girişken hem de enerjik bir yapıya sahip bireyleri ifade eder. Bu kişilik türü, cesur, hareketli ve iyimserdir ve motivasyonları kolaylıkla bozulmaz. İki renk de atak ve güçlü olduğundan, bu kombinasyona sahip bireyler etkileyici liderlik becerilerine sahiptir.

Kırmızı-sarı kişilikli bireyler, genellikle çelişki yaşamazlar ve sorunlara hızlı ve yaratıcı çözümler bulma konusunda yeteneklidirler. Buldukları sıra dışı çözümler, çevrelerindeki insanları şaşırtabilir ve etkileyebilir. Bu bireyler, güç odaklı hareket etme özelliğine sahipken, aynı zamanda eğlenmeyi de bilirler.

Ancak, bu tür bir kişilik kombinasyonuyla bazı zorluklar da yaşanabilir. Hiperaktif ve aşırı enerjik olmaları nedeniyle, bu bireylerin riskleri göz önünde bulundurarak hareket etmeleri ve desteklenmeleri önemlidir. Ayrıca, bu gruptaki bireylerin zaman zaman dikkat dağıtıcı ve düşüncesiz davranabileceği unutulmamalıdır.

Kırmızı-sarı kişilikli bireylerin başarılı olabilmeleri için, enerjilerini ve güçlerini doğru şekilde kullanmaları ve uygun dengeyi sağlamaları gerekmektedir. Özellikle ekip çalışmalarında ve liderlik pozisyonlarında, bu kişilik kombinasyonuyla sahip bireyler, etkili ve başarılı sonuçlar elde edebilirler. İş ve sosyal yaşamlarında, bu güçlü ve enerjik kişiliklerin, başkalarıyla uyum içinde çalışarak harika şeyler başarmaları mümkündür.
"""

mavi_kirmizi = """
Mavi ve kırmızı kişilik kombinasyonu, iş odaklı ve etkili bir karışımı temsil eder. Bu bireyler, kırmızının güçlü, kararlı ve sonuç odaklı yönleri ile mavinin planlı, programlı ve düzenli yönlerini birleştirir. Bu özellikler, bu tür kişiliklere sahip olan bireylerin işlerini önemseyerek, doğru ve hızlı bir şekilde sonuç almaya odaklanmalarına yardımcı olur.

Mavi-kırmızı kişilikli bireyler, sıklıkla iş ve sorumluluklarını çok ciddiye alır ve başarıları için sıkı çalışırlar. İki iş odaklı kişiliğin birleşmesi, bireyin zaman zaman yorulmasına ve stresli hissetmesine neden olabilir. Bununla birlikte, bu kişilik kombinasyonuna sahip olanlar genellikle hedeflerine ulaşmak için büyük çaba gösterir ve başarılı olma şansları yüksektir.

Ancak, bu kombinasyonda bazı zorluklar da bulunmaktadır. Özellikle, kırmızı ve mavi arasındaki dengeyi sağlamak ve her iki tarafın da güçlü yönlerini kullanmak önemlidir. Kırmızı, hızlı ve etkili kararlar almayı isterken, mavi daha planlı ve düşünceli hareket etmeyi tercih eder. Bu nedenle, mavi-kırmızı kişilikli bireyler, uyumlu bir şekilde çalışabilmek için bu iki yaklaşımı dengede tutmaya çalışmalıdır.

İş yaşamında ve kişisel ilişkilerinde, mavi-kırmızı kişilikli bireyler, kararlılık, düzen ve planlama becerileri sayesinde başarılı ve güvenilir olarak görülürler. İş birliği ve ekip çalışması gerektiren ortamlarda, bu kişilik kombinasyonuna sahip olanlar, etkili ve uyumlu bir şekilde çalışarak başarıya ulaşabilirler.
"""

yesil_sari = """
Yeşil ve sarı kişilik kombinasyonu, insan odaklı ve sosyal bir karışımı temsil eder. Bu bireyler, yeşilin empatik, barışçıl ve iş birliğine açık yönleri ile sarının enerjik, girişken ve eğlenceli yönlerini birleştirir. Bu özellikler, yeşil-sarı kişilikli bireylerin insanlarla kolayca iletişim kurmalarını ve sosyal ortamlarda sevilen, popüler bireyler olmalarını sağlar.

Yeşil-sarı kişilikli bireyler, sosyal ilişkilerinde başarılıdırlar ve genellikle insanlarla iyi anlaşır, hoş vakit geçirirler. İkna kabiliyetleri yüksek olduğu için, başkalarını kendi düşüncelerine ve fikirlerine çekmekte başarılı olurlar. Aynı zamanda, neşeli ve enerjik yapılarıyla ortamları canlandırabilir ve insanların moralini yükseltebilirler.

Ancak, bu kişilik kombinasyonunun dezavantajları da bulunmaktadır. Özellikle, yeşil ve sarı arasındaki dengeyi sağlamak ve her iki tarafın da güçlü yönlerini kullanmak önemlidir. Yeşil-sarı kişilikli bireyler, işleriyle ilgili olarak daha fazla odaklanma ve disipline ihtiyaç duyabilirler. İletişim odaklı oldukları için, zaman zaman işlerine konsantre olmakta ve önceliklerini belirlemekte zorlanabilirler.

İş yaşamında ve kişisel ilişkilerinde, yeşil-sarı kişilikli bireyler, insanlarla kurdukları güçlü bağlar ve empati yetenekleri sayesinde başarılı ve güvenilir olarak görülürler. İş birliği ve ekip çalışması gerektiren ortamlarda, bu kişilik kombinasyonuna sahip olanlar, insan ilişkilerini kullanarak etkili ve uyumlu bir şekilde çalışarak başarıya ulaşabilirler.
"""

mavi_yesil = """
Mavi ve yeşil kişilik kombinasyonu, temkinli ve içe dönük bir karışımı temsil eder. Bu bireyler, mavinin düzenli, planlı ve ciddi yönleri ile yeşilin barışçıl, sabırlı ve iş birliğine açık yönlerini birleştirir. Bu özellikler, mavi-yeşil kişilikli bireylerin sorumluluklarını titizlikle yerine getirmelerini ve dikkatli, düşünceli kararlar alarak hareket etmelerini sağlar.

Mavi-yeşil kişilikli bireyler, iş yaşamında başarılı olma potansiyeline sahiptirler. Detaylara dikkat etme ve planlama yetenekleri, onların projeleri başarıyla tamamlamalarına yardımcı olur. Aynı zamanda, yeşil tarafının getirdiği iş birliği ve uyum becerileri sayesinde, ekip çalışmasına uyum sağlayarak başkalarıyla verimli bir şekilde çalışabilirler.

Ancak, bu kişilik kombinasyonunun bazı zorlukları da vardır. Özellikle, mavi ve yeşil arasındaki dengeyi sağlamak önemlidir. Mavi-yeşil kişilikli bireyler, içe dönük yapıları nedeniyle sosyal ortamlarda zorlanabilir ve daha fazla iletişim becerisi geliştirmeye ihtiyaç duyabilirler. Ayrıca, karar verme süreçlerinde, mavi tarafının aşırı kuralcı ve analizci yaklaşımı ile yeşil tarafının riskten kaçınma eğilimi arasında denge kurmak önemlidir.

İş yaşamında ve kişisel ilişkilerinde, mavi-yeşil kişilikli bireyler, güvenilir, sadık ve destekleyici olarak görülürler. Başkalarına karşı anlayışlı ve empatik olma eğilimindedirler ve bu özellikleri sayesinde, dostluklar ve profesyonel ilişkiler kurarak başarıya ulaşabilirler.
"""

kirmizi_yesil = """
Kırmızı ve yeşil kişilik kombinasyonu, zıt karakter özelliklerine sahip renklerin bir araya gelmesiyle oluşan bir karışımdır. Kırmızı tarafı enerjik, dışa dönük, kararlı ve sonuç odaklıyken, yeşil tarafı daha içe dönük, kararsız ve risk almaktan çekinen yönleri ile öne çıkar. Bu kişilik karışımına sahip bireyler, hayatlarında sürekli iniş çıkışlar yaşayabilirler ve zaman zaman çelişkili davranışlar sergileyebilirler.

Kırmızı-yeşil kişilikli bireylerin güçlü yönleri arasında, kırmızının hızlı karar alma ve harekete geçme kabiliyeti ile yeşilin empati ve iş birliği becerileri bulunur. Bu özellikler sayesinde, bu bireyler hem liderlik hem de takım çalışması rollerinde başarılı olabilirler. Ayrıca, çeviklik ve uyum yetenekleri, değişen durumlara hızla uyum sağlamalarına yardımcı olur.

Ancak, bu kişilik kombinasyonunun bazı zorlukları da vardır. Özellikle, kırmızı ve yeşil arasındaki dengeyi sağlamak önemlidir. Kırmızı-yeşil kişilikli bireyler, bir yandan aldıkları kararları uygulamak isterken, diğer yandan bu kararları ertelemek isteyebilirler. Bu nedenle, bu bireyler için kararlılık ve istikrarı sağlamak zaman zaman zor olabilir.

İş yaşamında ve kişisel ilişkilerinde, kırmızı-yeşil kişilikli bireyler, güçlü iletişim becerileri ve sosyal yatkınlıkları sayesinde başkalarıyla kolaylıkla bağlantı kurabilirler. Fakat aynı zamanda, daha dengeli bir yaşam tarzı geliştirmek ve tutarlı davranışlar sergilemek için kendi içlerindeki çelişkilerle başa çıkmaları gerekebilir. Bu dengeyi sağlamaları durumunda hem iş hem de özel yaşamlarında başarı ve mutluluk yakalayabilirler.
"""

sari_mavi = """
Sarı ve mavi kişilik kombinasyonu, zıt karakter özelliklerine sahip renklerin bir araya gelmesiyle oluşan ilginç ve çelişkili bir karışımdır. Sarı tarafı eğlenceli, düzensiz, coşkulu ve enerjikken, mavi tarafı düzenli, planlı, ciddi ve detaylara önem veren yönleri ile öne çıkar. Bu kişilik karışımına sahip bireyler, hayatlarında zaman zaman dengesiz ve uyumsuz davranışlar sergileyebilirler.

Sarı-mavi kişilikli bireylerin güçlü yönleri arasında, sarının yaratıcılık ve iletişim becerileri ile mavinin analitik düşünme ve planlama kabiliyetleri bulunur. Bu özellikler sayesinde, bu bireyler hem yenilikçi fikirler üretebilir hem de bu fikirleri gerçekleştirmek için gerekli stratejileri geliştirebilirler. Ayrıca, sarı-mavi kişilikli bireyler, enerjik ve sosyal yönleri sayesinde insanlarla kolaylıkla iletişim kurabilir ve başkalarını etkileyebilirler.

Ancak, bu kişilik kombinasyonunun bazı zorlukları da vardır. Özellikle, sarı ve mavi arasındaki dengeyi sağlamak önemlidir. Sarı-mavi kişilikli bireyler, bir yandan eğlenceli ve spontane yaşamak isterken, diğer yandan düzenli ve planlı bir yaşam sürdürmek isteyebilirler. Bu nedenle, bu bireyler için yaşamlarında uyum ve dengeyi sağlamak zaman zaman zor olabilir.

İş yaşamında ve kişisel ilişkilerinde, sarı-mavi kişilikli bireyler, çelişkili doğalarına rağmen başarılı olabilirler. Bunun için, kendi içlerindeki zıt yönleri dengelemeyi öğrenmeleri ve her durumda en uygun özelliklerini kullanmayı başarmaları gerekmektedir. Bu dengeyi sağlamaları durumunda hem iş hem de özel yaşamlarında başarı ve mutluluk yakalayabilirler.
"""

nihai_sonuc_dict = {
    "kirmizi": "Kırmızı",
    "yesil": "Yeşil",
    "sari": "Sarı",
    "mavi": "Mavi",
    "kirmizi_sari": "Kırmızı-Sarı",
    "mavi_kirmizi": "Mavi-Kırmızı",
    "yesil_sari": "Yeşil-Sarı",
    "mavi_yesil": "Mavi-Yeşil",
    "kirmizi_yesil": "Kırmızı-Yeşil",
    "sari_mavi": "Sarı-Mavi",
}

nihai_sonuc_aciklama_dict = {
    "kirmizi": kirmizi,
    "yesil": yesil,
    "sari": sari,
    "mavi": mavi,
    "kirmizi_sari": kirmizi_sari,
    "mavi_kirmizi": mavi_kirmizi,
    "yesil_sari": yesil_sari,
    "mavi_yesil": mavi_yesil,
    "kirmizi_yesil": kirmizi_yesil,
    "sari_mavi": sari_mavi,
}


# Global değerler.
options = [0, 1, 2]
value_a1, value_a2, value_a3, value_a4 = 0, 0, 0, 0
value_b1, value_b2, value_b3, value_b4 = 0, 0, 0, 0
value_c1, value_c2, value_c3, value_c4 = 0, 0, 0, 0
value_d1, value_d2, value_d3, value_d4 = 0, 0, 0, 0
value_e1, value_e2, value_e3, value_e4 = 0, 0, 0, 0
value_f1, value_f2, value_f3, value_f4 = 0, 0, 0, 0
value_g1, value_g2, value_g3, value_g4 = 0, 0, 0, 0
value_h1, value_h2, value_h3, value_h4 = 0, 0, 0, 0
value_i1, value_i2, value_i3, value_i4 = 0, 0, 0, 0
value_j1, value_j2, value_j3, value_j4 = 0, 0, 0, 0
value_k1, value_k2, value_k3, value_k4 = 0, 0, 0, 0
value_m1, value_m2, value_m3, value_m4 = 0, 0, 0, 0
value_n1, value_n2, value_n3, value_n4 = 0, 0, 0, 0
value_o1, value_o2, value_o3, value_o4 = 0, 0, 0, 0
value_p1, value_p2, value_p3, value_p4 = 0, 0, 0, 0
sorunlu_cevaplar = []
ad_soyad, meslek, unvan, bolum, yas, cinsiyet, mail, telefon = (
    "",
    "",
    "",
    "",
    "",
    "",
    "",
    "",
)
erid = ""
resimle_ilgili_not = ""
dosya_adi_danisman = ""
dosya_adi_ogrenci = ""
nihai_sonuc = ""
nihai_sonuc_aciklama = ""


# Cevapları kontrol et
def cevaplari_kontrol_et(bir, iki, uc, dort):
    bir, iki, uc, dort = int(bir), int(iki), int(uc), int(dort)
    if (
        bir + iki + uc + dort == 3
        and ((bir + 1) * (iki + 1) * (uc + 1) * (dort + 1)) == 6
    ):
        pass
    elif (
        bir + iki + uc + dort == 2
        and ((bir + 1) * (iki + 1) * (uc + 1) * (dort + 1)) == 3
    ):
        pass
    elif bir + iki + uc + dort == 0:
        pass
    else:
        st.error(
            "Soruları cevaplama konusunda bir sorun var. "
            "Cevaplarınızda sadece bir adet '2' olmalıdır. "
            "Bunun yanında size yakın bir diğer seçeneğe '1' verebilirsiniz."
        )


# Hatalı cevapları tespit et
def hatali_mi(bir, iki, uc, dort):
    bir, iki, uc, dort = int(bir), int(iki), int(uc), int(dort)
    if (
        bir + iki + uc + dort == 3
        and ((bir + 1) * (iki + 1) * (uc + 1) * (dort + 1)) == 6
    ):
        pass
    elif (
        bir + iki + uc + dort == 2
        and ((bir + 1) * (iki + 1) * (uc + 1) * (dort + 1)) == 3
    ):
        pass
    elif bir + iki + uc + dort == 0:
        pass
    else:
        return "Hatalı"


# Hatalı cevapları kontrol et
def hatali_cevaplari_kontrol_et():
    global sorunlu_cevaplar
    sorunlu_cevaplar = []
    if hatali_mi(value_a1, value_a2, value_a3, value_a4) == "Hatalı":
        sorunlu_cevaplar.append("1. Soru")
    if hatali_mi(value_b1, value_b2, value_b3, value_b4) == "Hatalı":
        sorunlu_cevaplar.append("2. Soru")
    if hatali_mi(value_c1, value_c2, value_c3, value_c4) == "Hatalı":
        sorunlu_cevaplar.append("3. Soru")
    if hatali_mi(value_d1, value_d2, value_d3, value_d4) == "Hatalı":
        sorunlu_cevaplar.append("4. Soru")
    if hatali_mi(value_e1, value_e2, value_e3, value_e4) == "Hatalı":
        sorunlu_cevaplar.append("5. Soru")
    if hatali_mi(value_f1, value_f2, value_f3, value_f4) == "Hatalı":
        sorunlu_cevaplar.append("6. Soru")
    if hatali_mi(value_g1, value_g2, value_g3, value_g4) == "Hatalı":
        sorunlu_cevaplar.append("7. Soru")
    if hatali_mi(value_h1, value_h2, value_h3, value_h4) == "Hatalı":
        sorunlu_cevaplar.append("8. Soru")
    if hatali_mi(value_i1, value_i2, value_i3, value_i4) == "Hatalı":
        sorunlu_cevaplar.append("9. Soru")
    if hatali_mi(value_j1, value_j2, value_j3, value_j4) == "Hatalı":
        sorunlu_cevaplar.append("10. Soru")
    if hatali_mi(value_k1, value_k2, value_k3, value_k4) == "Hatalı":
        sorunlu_cevaplar.append("11. Soru")
    if hatali_mi(value_m1, value_m2, value_m3, value_m4) == "Hatalı":
        sorunlu_cevaplar.append("12. Soru")
    if hatali_mi(value_n1, value_n2, value_n3, value_n4) == "Hatalı":
        sorunlu_cevaplar.append("13. Soru")
    if hatali_mi(value_o1, value_o2, value_o3, value_o4) == "Hatalı":
        sorunlu_cevaplar.append("14. Soru")
    if hatali_mi(value_p1, value_p2, value_p3, value_p4) == "Hatalı":
        sorunlu_cevaplar.append("15. Soru")

    if len(sorunlu_cevaplar) > 0:
        st.error(
            f"Hatalı Cevaplar: {' '.join(sorunlu_cevaplar)} \n\nBu soruları tekrar gözden geçiriniz. "
            "Her sorunun altında dört şık yer almaktadır. Seçeneklerden sadece "
            "ikisine puan verebilirsiniz. Size en uygun seçeneğe 2, bir sonrakine "
            "ise 1 puan verin. Sadece tek seçenek seçilecek ise 2 puan verin. "
            "Lütfen bu kurallara uygun cevaplar veriniz."
        )
        return True
    else:
        return False


def head():
    st.markdown(
        """
        <h1 style='text-align: center'>
        Renkler ve Kişilik Envanteri
        </h1>
    """,
        unsafe_allow_html=True,
    )

    st.caption(
        """
        <p style='text-align: center;margin-bottom: 50px;'>
        Sakarya Üniversitesi Kariyer Geliştirme Koordinatörlüğü
        </p>
    """,
        unsafe_allow_html=True,
    )


def kisisel_bilgiler():
    global ad_soyad, meslek, unvan, bolum, yas, cinsiyet, mail, telefon
    ad_soyad = st.text_input("Adınız ve Soyadınız", key="ad_soyad")
    meslek = st.text_input("Mesleğiniz", key="meslek")
    unvan = st.text_input("Ünvanınız", key="unvan")
    bolum = st.text_input("Bölümünüz", key="bolum")
    yas = st.text_input("Yaşınız", key="yas")
    cinsiyet = st.selectbox("Cinsiyetiniz", ["Erkek", "Kadın"], key="cinsiyet")
    mail = st.text_input("Mail Adresiniz", key="mail")
    telefon = st.text_input("Telefon Numaranız", key="telefon")


def sorular():
    global value_a1, value_a2, value_a3, value_a4, value_b1, value_b2, value_b3, value_b4, value_c1, value_c2, value_c3, value_c4, value_d1, value_d2, value_d3, value_d4, value_e1, value_e2, value_e3, value_e4, value_f1, value_f2, value_f3, value_f4, value_g1, value_g2, value_g3, value_g4, value_h1, value_h2, value_h3, value_h4, value_i1, value_i2, value_i3, value_i4, value_j1, value_j2, value_j3, value_j4, value_k1, value_k2, value_k3, value_k4, value_m1, value_m2, value_m3, value_m4, value_n1, value_n2, value_n3, value_n4, value_o1, value_o2, value_o3, value_o4, value_p1, value_p2, value_p3, value_p4
    st.info(
        "Farklı karakterlerin özelliklerini tespit etmede bir ölçü olarak aşağıda bir test uygulaması "
        "yer almaktadır. Aşağıdaki test, her biri dört şıktan oluşan 15 soru içermektedir. Seçeneklerde belirtilen "
        "özellikler, size ne derece uyar ise ona göre puan veriniz."
    )
    st.warning(
        "Her sorunun altında dört şık yer almaktadır. Seçeneklerden sadece "
        "ikisine puan verebilirsiniz. Size en uygun seçeneğe 2, bir sonrakine "
        "ise 1 puan verin. Sadece tek seçenek seçilecek ise 2 puan verin."
    )
    st.success(
        "Envanteri telefondan dolduruyorsanız, telefonu yatay modda kullanmak okunabilirliği arttırabilir."
    )
    with st.container():
        st.markdown("<div class='content-container'></div>", unsafe_allow_html=True)

        st.subheader(a)
        st.markdown("<div class='row-divider'></div>", unsafe_allow_html=True)

        col1, col2 = st.columns([9, 1])
        with col1.container():
            col1.markdown(f"{a1}", unsafe_allow_html=True)
            value_a1 = col2.selectbox("Seçin", options, key="a1")

        st.markdown("<div class='row-divider'></div>", unsafe_allow_html=True)

        col1, col2 = st.columns([9, 1])
        with col1.container():
            col1.markdown(f"{a2}", unsafe_allow_html=True)
            value_a2 = col2.selectbox("Seçin", options, key="a2")

        st.markdown("<div class='row-divider'></div>", unsafe_allow_html=True)

        col1, col2 = st.columns([9, 1])
        with col1.container():
            col1.markdown(f"{a3}", unsafe_allow_html=True)
            value_a3 = col2.selectbox("Seçin", options, key="a3")

        st.markdown("<div class='row-divider'></div>", unsafe_allow_html=True)

        col1, col2 = st.columns([9, 1])
        with col1.container():
            col1.markdown(f"{a4}", unsafe_allow_html=True)
            value_a4 = col2.selectbox("Seçin", options, key="a4")

        cevaplari_kontrol_et(value_a1, value_a2, value_a3, value_a4)
        st.markdown("<div class='content-container'></div>", unsafe_allow_html=True)

    with st.container():
        st.subheader(b)

        st.markdown("<div class='row-divider'></div>", unsafe_allow_html=True)

        col1, col2 = st.columns([9, 1])
        with col1.container():
            col1.markdown(f"{b1}", unsafe_allow_html=True)
            value_b1 = col2.selectbox("Seçin", options, key="b1")

        st.markdown("<div class='row-divider'></div>", unsafe_allow_html=True)

        col1, col2 = st.columns([9, 1])
        with col1.container():
            col1.markdown(f"{b2}", unsafe_allow_html=True)
            value_b2 = col2.selectbox("Seçin", options, key="b2")

        st.markdown("<div class='row-divider'></div>", unsafe_allow_html=True)

        col1, col2 = st.columns([9, 1])
        with col1.container():
            col1.markdown(f"{b3}", unsafe_allow_html=True)
            value_b3 = col2.selectbox("Seçin", options, key="b3")

        st.markdown("<div class='row-divider'></div>", unsafe_allow_html=True)

        col1, col2 = st.columns([9, 1])
        with col1.container():
            col1.markdown(f"{b4}", unsafe_allow_html=True)
            value_b4 = col2.selectbox("Seçin", options, key="b4")

        cevaplari_kontrol_et(value_b1, value_b2, value_b3, value_b4)
        st.markdown("<div class='content-container'></div>", unsafe_allow_html=True)

    with st.container():
        st.subheader(c)

        st.markdown("<div class='row-divider'></div>", unsafe_allow_html=True)

        col1, col2 = st.columns([9, 1])
        with col1.container():
            col1.markdown(f"{c1}", unsafe_allow_html=True)
            value_c1 = col2.selectbox("Seçin", options, key="c1")

        st.markdown("<div class='row-divider'></div>", unsafe_allow_html=True)

        col1, col2 = st.columns([9, 1])
        with col1.container():
            col1.markdown(f"{c2}", unsafe_allow_html=True)
            value_c2 = col2.selectbox("Seçin", options, key="c2")

        st.markdown("<div class='row-divider'></div>", unsafe_allow_html=True)

        col1, col2 = st.columns([9, 1])
        with col1.container():
            col1.markdown(f"{c3}", unsafe_allow_html=True)
            value_c3 = col2.selectbox("Seçin", options, key="c3")

        st.markdown("<div class='row-divider'></div>", unsafe_allow_html=True)

        col1, col2 = st.columns([9, 1])
        with col1.container():
            col1.markdown(f"{c4}", unsafe_allow_html=True)
            value_c4 = col2.selectbox("Seçin", options, key="c4")

        cevaplari_kontrol_et(value_c1, value_c2, value_c3, value_c4)
        st.markdown("<div class='content-container'></div>", unsafe_allow_html=True)

    with st.container():
        st.subheader(d)

        st.markdown("<div class='row-divider'></div>", unsafe_allow_html=True)

        col1, col2 = st.columns([9, 1])
        with col1.container():
            col1.markdown(f"{d1}", unsafe_allow_html=True)
            value_d1 = col2.selectbox("Seçin", options, key="d1")

        st.markdown("<div class='row-divider'></div>", unsafe_allow_html=True)

        col1, col2 = st.columns([9, 1])
        with col1.container():
            col1.markdown(f"{d2}", unsafe_allow_html=True)
            value_d2 = col2.selectbox("Seçin", options, key="d2")

        st.markdown("<div class='row-divider'></div>", unsafe_allow_html=True)

        col1, col2 = st.columns([9, 1])
        with col1.container():
            col1.markdown(f"{d3}", unsafe_allow_html=True)
            value_d3 = col2.selectbox("Seçin", options, key="d3")

        st.markdown("<div class='row-divider'></div>", unsafe_allow_html=True)

        col1, col2 = st.columns([9, 1])
        with col1.container():
            col1.markdown(f"{d4}", unsafe_allow_html=True)
            value_d4 = col2.selectbox("Seçin", options, key="d4")

        cevaplari_kontrol_et(value_d1, value_d2, value_d3, value_d4)
        st.markdown("<div class='content-container'></div>", unsafe_allow_html=True)

    with st.container():
        st.subheader(e)

        st.markdown("<div class='row-divider'></div>", unsafe_allow_html=True)

        col1, col2 = st.columns([9, 1])
        with col1.container():
            col1.markdown(f"{e1}", unsafe_allow_html=True)
            value_e1 = col2.selectbox("Seçin", options, key="e1")

        st.markdown("<div class='row-divider'></div>", unsafe_allow_html=True)

        col1, col2 = st.columns([9, 1])
        with col1.container():
            col1.markdown(f"{e2}", unsafe_allow_html=True)
            value_e2 = col2.selectbox("Seçin", options, key="e2")

        st.markdown("<div class='row-divider'></div>", unsafe_allow_html=True)

        col1, col2 = st.columns([9, 1])
        with col1.container():
            col1.markdown(f"{e3}", unsafe_allow_html=True)
            value_e3 = col2.selectbox("Seçin", options, key="e3")

        st.markdown("<div class='row-divider'></div>", unsafe_allow_html=True)

        col1, col2 = st.columns([9, 1])
        with col1.container():
            col1.markdown(f"{e4}", unsafe_allow_html=True)
            value_e4 = col2.selectbox("Seçin", options, key="e4")

        cevaplari_kontrol_et(value_e1, value_e2, value_e3, value_e4)
        st.markdown("<div class='content-container'></div>", unsafe_allow_html=True)

    with st.container():
        st.subheader(f)

        st.markdown("<div class='row-divider'></div>", unsafe_allow_html=True)

        col1, col2 = st.columns([9, 1])
        with col1.container():
            col1.markdown(f"{f1}", unsafe_allow_html=True)
            value_f1 = col2.selectbox("Seçin", options, key="f1")

        st.markdown("<div class='row-divider'></div>", unsafe_allow_html=True)

        col1, col2 = st.columns([9, 1])
        with col1.container():
            col1.markdown(f"{f2}", unsafe_allow_html=True)
            value_f2 = col2.selectbox("Seçin", options, key="f2")

        st.markdown("<div class='row-divider'></div>", unsafe_allow_html=True)

        col1, col2 = st.columns([9, 1])
        with col1.container():
            col1.markdown(f"{f3}", unsafe_allow_html=True)
            value_f3 = col2.selectbox("Seçin", options, key="f3")

        st.markdown("<div class='row-divider'></div>", unsafe_allow_html=True)

        col1, col2 = st.columns([9, 1])
        with col1.container():
            col1.markdown(f"{f4}", unsafe_allow_html=True)
            value_f4 = col2.selectbox("Seçin", options, key="f4")

        cevaplari_kontrol_et(value_f1, value_f2, value_f3, value_f4)
        st.markdown("<div class='content-container'></div>", unsafe_allow_html=True)

    with st.container():
        st.subheader(g)

        st.markdown("<div class='row-divider'></div>", unsafe_allow_html=True)

        col1, col2 = st.columns([9, 1])
        with col1.container():
            col1.markdown(f"{g1}", unsafe_allow_html=True)
            value_g1 = col2.selectbox("Seçin", options, key="g1")

        st.markdown("<div class='row-divider'></div>", unsafe_allow_html=True)

        col1, col2 = st.columns([9, 1])
        with col1.container():
            col1.markdown(f"{g2}", unsafe_allow_html=True)
            value_g2 = col2.selectbox("Seçin", options, key="g2")

        st.markdown("<div class='row-divider'></div>", unsafe_allow_html=True)

        col1, col2 = st.columns([9, 1])
        with col1.container():
            col1.markdown(f"{g3}", unsafe_allow_html=True)
            value_g3 = col2.selectbox("Seçin", options, key="g3")

        st.markdown("<div class='row-divider'></div>", unsafe_allow_html=True)

        col1, col2 = st.columns([9, 1])
        with col1.container():
            col1.markdown(f"{g4}", unsafe_allow_html=True)
            value_g4 = col2.selectbox("Seçin", options, key="g4")

        cevaplari_kontrol_et(value_g1, value_g2, value_g3, value_g4)
        st.markdown("<div class='content-container'></div>", unsafe_allow_html=True)

    with st.container():
        st.subheader(h)

        st.markdown("<div class='row-divider'></div>", unsafe_allow_html=True)

        col1, col2 = st.columns([9, 1])
        with col1.container():
            col1.markdown(f"{h1}", unsafe_allow_html=True)
            value_h1 = col2.selectbox("Seçin", options, key="h1")

        st.markdown("<div class='row-divider'></div>", unsafe_allow_html=True)

        col1, col2 = st.columns([9, 1])
        with col1.container():
            col1.markdown(f"{h2}", unsafe_allow_html=True)
            value_h2 = col2.selectbox("Seçin", options, key="h2")

        st.markdown("<div class='row-divider'></div>", unsafe_allow_html=True)

        col1, col2 = st.columns([9, 1])
        with col1.container():
            col1.markdown(f"{h3}", unsafe_allow_html=True)
            value_h3 = col2.selectbox("Seçin", options, key="h3")

        st.markdown("<div class='row-divider'></div>", unsafe_allow_html=True)

        col1, col2 = st.columns([9, 1])
        with col1.container():
            col1.markdown(f"{h4}", unsafe_allow_html=True)
            value_h4 = col2.selectbox("Seçin", options, key="h4")

        cevaplari_kontrol_et(value_h1, value_h2, value_h3, value_h4)
        st.markdown("<div class='content-container'></div>", unsafe_allow_html=True)

    with st.container():
        st.subheader(i)

        st.markdown("<div class='row-divider'></div>", unsafe_allow_html=True)

        col1, col2 = st.columns([9, 1])
        with col1.container():
            col1.markdown(f"{i1}", unsafe_allow_html=True)
            value_i1 = col2.selectbox("Seçin", options, key="i1")

        st.markdown("<div class='row-divider'></div>", unsafe_allow_html=True)

        col1, col2 = st.columns([9, 1])
        with col1.container():
            col1.markdown(f"{i2}", unsafe_allow_html=True)
            value_i2 = col2.selectbox("Seçin", options, key="i2")

        st.markdown("<div class='row-divider'></div>", unsafe_allow_html=True)

        col1, col2 = st.columns([9, 1])
        with col1.container():
            col1.markdown(f"{i3}", unsafe_allow_html=True)
            value_i3 = col2.selectbox("Seçin", options, key="i3")

        st.markdown("<div class='row-divider'></div>", unsafe_allow_html=True)

        col1, col2 = st.columns([9, 1])
        with col1.container():
            col1.markdown(f"{i4}", unsafe_allow_html=True)
            value_i4 = col2.selectbox("Seçin", options, key="i4")

        cevaplari_kontrol_et(value_i1, value_i2, value_i3, value_i4)
        st.markdown("<div class='content-container'></div>", unsafe_allow_html=True)

    with st.container():
        st.subheader(j)

        st.markdown("<div class='row-divider'></div>", unsafe_allow_html=True)

        col1, col2 = st.columns([9, 1])
        with col1.container():
            col1.markdown(f"{j1}", unsafe_allow_html=True)
            value_j1 = col2.selectbox("Seçin", options, key="j1")

        st.markdown("<div class='row-divider'></div>", unsafe_allow_html=True)

        col1, col2 = st.columns([9, 1])
        with col1.container():
            col1.markdown(f"{j2}", unsafe_allow_html=True)
            value_j2 = col2.selectbox("Seçin", options, key="j2")

        st.markdown("<div class='row-divider'></div>", unsafe_allow_html=True)

        col1, col2 = st.columns([9, 1])
        with col1.container():
            col1.markdown(f"{j3}", unsafe_allow_html=True)
            value_j3 = col2.selectbox("Seçin", options, key="j3")

        st.markdown("<div class='row-divider'></div>", unsafe_allow_html=True)

        col1, col2 = st.columns([9, 1])
        with col1.container():
            col1.markdown(f"{j4}", unsafe_allow_html=True)
            value_j4 = col2.selectbox("Seçin", options, key="j4")

        cevaplari_kontrol_et(value_j1, value_j2, value_j3, value_j4)
        st.markdown("<div class='content-container'></div>", unsafe_allow_html=True)

    with st.container():
        st.subheader(k)

        st.markdown("<div class='row-divider'></div>", unsafe_allow_html=True)

        col1, col2 = st.columns([9, 1])
        with col1.container():
            col1.markdown(f"{k1}", unsafe_allow_html=True)
            value_k1 = col2.selectbox("Seçin", options, key="k1")

        st.markdown("<div class='row-divider'></div>", unsafe_allow_html=True)

        col1, col2 = st.columns([9, 1])
        with col1.container():
            col1.markdown(f"{k2}", unsafe_allow_html=True)
            value_k2 = col2.selectbox("Seçin", options, key="k2")

        st.markdown("<div class='row-divider'></div>", unsafe_allow_html=True)

        col1, col2 = st.columns([9, 1])
        with col1.container():
            col1.markdown(f"{k3}", unsafe_allow_html=True)
            value_k3 = col2.selectbox("Seçin", options, key="k3")

        st.markdown("<div class='row-divider'></div>", unsafe_allow_html=True)

        col1, col2 = st.columns([9, 1])
        with col1.container():
            col1.markdown(f"{k4}", unsafe_allow_html=True)
            value_k4 = col2.selectbox("Seçin", options, key="k4")

        cevaplari_kontrol_et(value_k1, value_k2, value_k3, value_k4)
        st.markdown("<div class='content-container'></div>", unsafe_allow_html=True)

    with st.container():
        st.subheader(m)

        st.markdown("<div class='row-divider'></div>", unsafe_allow_html=True)

        col1, col2 = st.columns([9, 1])
        with col1.container():
            col1.markdown(f"{m1}", unsafe_allow_html=True)
            value_m1 = col2.selectbox("Seçin", options, key="m1")

        st.markdown("<div class='row-divider'></div>", unsafe_allow_html=True)

        col1, col2 = st.columns([9, 1])
        with col1.container():
            col1.markdown(f"{m2}", unsafe_allow_html=True)
            value_m2 = col2.selectbox("Seçin", options, key="m2")

        st.markdown("<div class='row-divider'></div>", unsafe_allow_html=True)

        col1, col2 = st.columns([9, 1])
        with col1.container():
            col1.markdown(f"{m3}", unsafe_allow_html=True)
            value_m3 = col2.selectbox("Seçin", options, key="m3")

        st.markdown("<div class='row-divider'></div>", unsafe_allow_html=True)

        col1, col2 = st.columns([9, 1])
        with col1.container():
            col1.markdown(f"{m4}", unsafe_allow_html=True)
            value_m4 = col2.selectbox("Seçin", options, key="m4")

        cevaplari_kontrol_et(value_m1, value_m2, value_m3, value_m4)
        st.markdown("<div class='content-container'></div>", unsafe_allow_html=True)

    with st.container():
        st.subheader(n)

        st.markdown("<div class='row-divider'></div>", unsafe_allow_html=True)

        col1, col2 = st.columns([9, 1])
        with col1.container():
            col1.markdown(f"{n1}", unsafe_allow_html=True)
            value_n1 = col2.selectbox("Seçin", options, key="n1")

        st.markdown("<div class='row-divider'></div>", unsafe_allow_html=True)

        col1, col2 = st.columns([9, 1])
        with col1.container():
            col1.markdown(f"{n2}", unsafe_allow_html=True)
            value_n2 = col2.selectbox("Seçin", options, key="n2")

        st.markdown("<div class='row-divider'></div>", unsafe_allow_html=True)

        col1, col2 = st.columns([9, 1])
        with col1.container():
            col1.markdown(f"{n3}", unsafe_allow_html=True)
            value_n3 = col2.selectbox("Seçin", options, key="n3")

        st.markdown("<div class='row-divider'></div>", unsafe_allow_html=True)

        col1, col2 = st.columns([9, 1])
        with col1.container():
            col1.markdown(f"{n4}", unsafe_allow_html=True)
            value_n4 = col2.selectbox("Seçin", options, key="n4")

        cevaplari_kontrol_et(value_n1, value_n2, value_n3, value_n4)
        st.markdown("<div class='content-container'></div>", unsafe_allow_html=True)

    with st.container():
        st.subheader(o)

        st.markdown("<div class='row-divider'></div>", unsafe_allow_html=True)

        col1, col2 = st.columns([9, 1])
        with col1.container():
            col1.markdown(f"{o1}", unsafe_allow_html=True)
            value_o1 = col2.selectbox("Seçin", options, key="o1")

        st.markdown("<div class='row-divider'></div>", unsafe_allow_html=True)

        col1, col2 = st.columns([9, 1])
        with col1.container():
            col1.markdown(f"{o2}", unsafe_allow_html=True)
            value_o2 = col2.selectbox("Seçin", options, key="o2")

        st.markdown("<div class='row-divider'></div>", unsafe_allow_html=True)

        col1, col2 = st.columns([9, 1])
        with col1.container():
            col1.markdown(f"{o3}", unsafe_allow_html=True)
            value_o3 = col2.selectbox("Seçin", options, key="o3")

        st.markdown("<div class='row-divider'></div>", unsafe_allow_html=True)

        col1, col2 = st.columns([9, 1])
        with col1.container():
            col1.markdown(f"{o4}", unsafe_allow_html=True)
            value_o4 = col2.selectbox("Seçin", options, key="o4")

        cevaplari_kontrol_et(value_o1, value_o2, value_o3, value_o4)
        st.markdown("<div class='content-container'></div>", unsafe_allow_html=True)

    with st.container():
        st.subheader(p)

        st.markdown("<div class='row-divider'></div>", unsafe_allow_html=True)

        col1, col2 = st.columns([9, 1])
        with col1.container():
            col1.markdown(f"{p1}", unsafe_allow_html=True)
            value_p1 = col2.selectbox("Seçin", options, key="p1")

        st.markdown("<div class='row-divider'></div>", unsafe_allow_html=True)

        col1, col2 = st.columns([9, 1])
        with col1.container():
            col1.markdown(f"{p2}", unsafe_allow_html=True)
            value_p2 = col2.selectbox("Seçin", options, key="p2")

        st.markdown("<div class='row-divider'></div>", unsafe_allow_html=True)

        col1, col2 = st.columns([9, 1])
        with col1.container():
            col1.markdown(f"{p3}", unsafe_allow_html=True)
            value_p3 = col2.selectbox("Seçin", options, key="p3")

        st.markdown("<div class='row-divider'></div>", unsafe_allow_html=True)

        col1, col2 = st.columns([9, 1])
        with col1.container():
            col1.markdown(f"{p4}", unsafe_allow_html=True)
            value_p4 = col2.selectbox("Seçin", options, key="p4")

        cevaplari_kontrol_et(value_p1, value_p2, value_p3, value_p4)
        st.markdown("<div class='content-container'></div>", unsafe_allow_html=True)


def dosya_olustur_danisman():
    global dosya_adi_danisman
    document = Document()

    section = document.sections[0]
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.text = (
        "\tSakarya Üniversitesi Kariyer Geliştirme Koordinatörlüğü"
    )
    paragraph.style = document.styles["Header"]

    document.add_heading(ad_soyad.title(), 0)
    p = document.add_paragraph()
    p.add_run("Doldurulma tarihi").bold = True
    p.add_run(
        f": {(datetime.now() + timedelta(hours=3)).strftime('%d/%m/%Y %H:%M:%S')}"
    ).italic = True

    document.add_heading("A. Kişisel Bilgiler", level=1)

    p = document.add_paragraph()
    p.add_run("Mesleği").bold = True
    p.add_run(f": {meslek}").italic = True

    p = document.add_paragraph()
    p.add_run("Ünvanı").bold = True
    p.add_run(f": {unvan}").italic = True

    p = document.add_paragraph()
    p.add_run("Bölümü").bold = True
    p.add_run(f": {bolum}").italic = True

    p = document.add_paragraph()
    p.add_run("Yaş").bold = True
    p.add_run(f": {yas}").italic = True

    p = document.add_paragraph()
    p.add_run("Cinsiyet").bold = True
    p.add_run(f": {cinsiyet}").italic = True

    p = document.add_paragraph()
    p.add_run("Mail").bold = True
    p.add_run(f": {mail}").italic = True

    document.add_heading("B. Analiz Sonuçları", level=1)

    gorseli_olustur()
    document.add_picture("gradient.png")

    nihai_karar()
    p = document.add_paragraph()
    p.add_run("Renginiz: ").bold = True
    p.add_run(nihai_sonuc).italic = True

    p = document.add_paragraph()
    p.add_run(nihai_sonuc_aciklama)

    erid = "".join(random.choices(string.ascii_uppercase + string.ascii_lowercase, k=5))

    p = document.add_paragraph()
    p.add_run(resimle_ilgili_not).italic = True

    p = document.add_paragraph()
    p.add_run("ERID: ").bold = True
    p.add_run(f"{erid}").italic = True

    # Kullanıcının verdiği bütün cevapları ekle.
    document.add_heading("C. Cevaplar", level=1)
    # Soru, sorunun cevapları ve kullanıcının verdiği cevapları ekle. Kompakt bir tasarım yap.
    sorular = [
        "a",
        "b",
        "c",
        "d",
        "e",
        "f",
        "g",
        "h",
        "i",
        "j",
        "k",
        "m",
        "n",
        "o",
        "p",
    ]
    for index, soru in enumerate(sorular):
        document.add_heading(f"{globals()[f'{soru}']}", level=2)
        document.add_paragraph(
            f"{globals()[f'{soru}1']} - Verilen Puan ({globals()[f'value_{soru}1']}) \n"
            f"{globals()[f'{soru}2']} - Verilen Puan ({globals()[f'value_{soru}2']}) \n"
            f"{globals()[f'{soru}3']} - Verilen Puan ({globals()[f'value_{soru}3']}) \n"
            f"{globals()[f'{soru}4']} - Verilen Puan ({globals()[f'value_{soru}4']}) \n"
        )

    dosya_adi_danisman = f"{ad_soyad.title()}-Danışman.docx"
    document.save(dosya_adi_danisman)


def dosya_olustur_ogrenci():
    global dosya_adi_ogrenci
    document = Document()

    section = document.sections[0]
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.text = (
        "\tSakarya Üniversitesi Kariyer Geliştirme Koordinatörlüğü"
    )
    paragraph.style = document.styles["Header"]

    document.add_heading(ad_soyad.title(), 0)
    p = document.add_paragraph()
    p.add_run("Doldurulma tarihi").bold = True
    p.add_run(
        f": {(datetime.now() + timedelta(hours=3)).strftime('%d/%m/%Y %H:%M:%S')}"
    ).italic = True

    document.add_heading("A. Kişisel Bilgiler", level=1)

    p = document.add_paragraph()
    p.add_run("Mesleği").bold = True
    p.add_run(f": {meslek}").italic = True

    p = document.add_paragraph()
    p.add_run("Ünvanı").bold = True
    p.add_run(f": {unvan}").italic = True

    p = document.add_paragraph()
    p.add_run("Bölümü").bold = True
    p.add_run(f": {bolum}").italic = True

    p = document.add_paragraph()
    p.add_run("Yaş").bold = True
    p.add_run(f": {yas}").italic = True

    p = document.add_paragraph()
    p.add_run("Cinsiyet").bold = True
    p.add_run(f": {cinsiyet}").italic = True

    p = document.add_paragraph()
    p.add_run("Mail").bold = True
    p.add_run(f": {mail}").italic = True

    document.add_heading("B. Analiz Sonuçları", level=1)

    gorseli_olustur()
    document.add_picture("gradient.png")

    nihai_karar()
    p = document.add_paragraph()
    p.add_run("Renginiz: ").bold = True
    p.add_run(nihai_sonuc).italic = True

    p = document.add_paragraph()
    p.add_run(nihai_sonuc_aciklama)

    erid = "".join(random.choices(string.ascii_uppercase + string.ascii_lowercase, k=5))

    p = document.add_paragraph()
    p.add_run(resimle_ilgili_not).italic = True

    p = document.add_paragraph()
    p.add_run("ERID: ").bold = True
    p.add_run(f"{erid}").italic = True

    dosya_adi_ogrenci = f"{ad_soyad.title()}.docx"
    document.save(dosya_adi_ogrenci)


def kontrol_butonu():
    global mail, dosya_adi_danisman, dosya_adi_ogrenci
    if st.button("Envanteri Kaydet"):
        # Bütün cevapları kontrol et, sorun olanları ekrana yazdır.
        if not hatali_cevaplari_kontrol_et():
            if mail == "":
                st.error("Mail adresi boş olamaz.")
            else:
                with st.empty():
                    st.success("Analiz oluşturuluyor...")
                    dosya_olustur_ogrenci()
                    dosya_olustur_danisman()
                    # st.success("Analiz oluşturuldu! Kaydediliyor...")
                    # dosyayi_kaydet()
                    send_mail.send_analysis(mail, [dosya_adi_ogrenci])
                    send_mail.send_analysis_to_danisman(
                        "cansusen@sakarya.edu.tr", [dosya_adi_danisman]
                    )

                    st.success(
                        "Analiz Kaydedildi ve mail adresinize gönderildi, kariyer@sakarya.edu.tr üzerinden bizimle "
                        "iletişime geçebilirsiniz."
                    )
                st.balloons()


def create_cred_file():
    cred_str = (
        '{"access_token": "'
        + st.secrets["access_token"]
        + '", "client_id": "'
        + st.secrets["client_id"]
        + '", "client_secret": "'
        + st.secrets["client_secret"]
        + '", "refresh_token": "'
        + st.secrets["refresh_token"]
        + '", "token_expiry": "'
        + st.secrets["token_expiry"]
        + '", "token_uri": "'
        + st.secrets["token_uri"]
        + '", "user_agent": null'
        + ', "revoke_uri": "'
        + st.secrets["revoke_uri"]
        + '", "id_token": null'
        + ', "id_token_jwt": null'
        + ', "token_response": {"access_token": "'
        + st.secrets["token_response"]["access_token"]
        + '", "expires_in": '
        + str(st.secrets["token_response"]["expires_in"])
        + ', "refresh_token": "'
        + st.secrets["token_response"]["refresh_token"]
        + '", "scope": "'
        + st.secrets["token_response"]["scope"]
        + '", "token_type": "'
        + st.secrets["token_response"]["token_type"]
        + '"}, "scopes": ["https://www.googleapis.com/auth/drive"]'
        + ', "token_info_uri": "'
        + st.secrets["token_info_uri"]
        + '", "invalid": false'
        + ', "_class": "'
        + st.secrets["_class"]
        + '", "_module": "'
        + st.secrets["_module"]
        + '"}'
    )
    text_file = open("mycreds_test.txt", "w")
    text_file.write(cred_str)
    text_file.close()


def gorseli_olustur():
    global resimle_ilgili_not
    total_1 = (
        int(value_a1)
        + int(value_b1)
        + int(value_c1)
        + int(value_d1)
        + int(value_e1)
        + int(value_f1)
        + int(value_g1)
        + int(value_h1)
        + int(value_i1)
        + int(value_j1)
        + int(value_k1)
        + int(value_m1)
        + int(value_n1)
        + int(value_o1)
        + int(value_p1)
    )
    total_2 = (
        int(value_a2)
        + int(value_b2)
        + int(value_c2)
        + int(value_d2)
        + int(value_e2)
        + int(value_f2)
        + int(value_g2)
        + int(value_h2)
        + int(value_i2)
        + int(value_j2)
        + int(value_k2)
        + int(value_m2)
        + int(value_n2)
        + int(value_o2)
        + int(value_p2)
    )
    total_3 = (
        int(value_a3)
        + int(value_b3)
        + int(value_c3)
        + int(value_d3)
        + int(value_e3)
        + int(value_f3)
        + int(value_g3)
        + int(value_h3)
        + int(value_i3)
        + int(value_j3)
        + int(value_k3)
        + int(value_m3)
        + int(value_n3)
        + int(value_o3)
        + int(value_p3)
    )
    total_4 = (
        int(value_a4)
        + int(value_b4)
        + int(value_c4)
        + int(value_d4)
        + int(value_e4)
        + int(value_f4)
        + int(value_g4)
        + int(value_h4)
        + int(value_i4)
        + int(value_j4)
        + int(value_k4)
        + int(value_m4)
        + int(value_n4)
        + int(value_o4)
        + int(value_p4)
    )

    total_names = ["kırmızı", "sarı", "yeşil", "mavi"]
    total = total_1 + total_2 + total_3 + total_4
    try:
        total1 = total_1 / total
    except ZeroDivisionError as zero_exception:
        print(zero_exception)
        total1 = 0
    try:
        total2 = total_2 / total
    except ZeroDivisionError as zero_exception:
        print(zero_exception)
        total2 = 0
    try:
        total3 = total_3 / total
    except ZeroDivisionError as zero_exception:
        print(zero_exception)
        total3 = 0
    try:
        total4 = total_4 / total
    except ZeroDivisionError as zero_exception:
        print(zero_exception)
        total4 = 0

    rational = [total1, total2, total3, total4]
    color_tone.create_gradient_image(total_names, rational, 430, 80, "gradient.png")

    color_distribution = zip(total_names, rational)
    color_distribution = sorted(color_distribution, key=lambda x: x[1], reverse=True)

    if 1 in rational:
        resimle_ilgili_not = (
            "* Resimle ilgili ufak bir not: Analiziniz sonucunda sadece bir renk çıktı. Ancak "
            "bu demek değildir ki sadece bir renk ifade ediyorsunuz. Hayatta her şey siyah veya "
            "beyaz olmadığı gibi, tek renk de değildir. Sadece bir renk çıksanız bile, bu rengin "
            "birçok tonu olduğunu unutmayın."
        )
    else:
        resimle_ilgili_not = (
            "* Resimle ilgili ufak bir not: Analiziniz sonucunda birkaç renk çıktı. Ancak "
            "bu demek değildir ki sadece birkaç renk ifade ediyorsunuz. Hayatta her şey siyah veya "
            "beyaz olmadığı gibi, sadece birkaç renk de değildir. Sadece birkaç renk çıksanız bile, "
            "bu renklerin birçok tonu olduğunu unutmayın. Bu görsel tamamen envanter "
            "sonucunuzdan elde edildi. Verdiğiniz cevaplar dolayısıyla tamamen size özeldir. ("
        )

        for k in color_distribution:
            if k[1] != 0:
                resimle_ilgili_not += f" %{k[1] * 100:.2f} {k[0]}"
        resimle_ilgili_not += ")"

    # Kullanıcıların verdiği cevaplar.


def nihai_karar():
    global nihai_sonuc, nihai_sonuc_aciklama

    total_1 = (
        int(value_a1)
        + int(value_b1)
        + int(value_c1)
        + int(value_d1)
        + int(value_e1)
        + int(value_f1)
        + int(value_g1)
        + int(value_h1)
        + int(value_i1)
        + int(value_j1)
        + int(value_k1)
        + int(value_m1)
        + int(value_n1)
        + int(value_o1)
        + int(value_p1)
    )
    total_2 = (
        int(value_a2)
        + int(value_b2)
        + int(value_c2)
        + int(value_d2)
        + int(value_e2)
        + int(value_f2)
        + int(value_g2)
        + int(value_h2)
        + int(value_i2)
        + int(value_j2)
        + int(value_k2)
        + int(value_m2)
        + int(value_n2)
        + int(value_o2)
        + int(value_p2)
    )
    total_3 = (
        int(value_a3)
        + int(value_b3)
        + int(value_c3)
        + int(value_d3)
        + int(value_e3)
        + int(value_f3)
        + int(value_g3)
        + int(value_h3)
        + int(value_i3)
        + int(value_j3)
        + int(value_k3)
        + int(value_m3)
        + int(value_n3)
        + int(value_o3)
        + int(value_p3)
    )
    total_4 = (
        int(value_a4)
        + int(value_b4)
        + int(value_c4)
        + int(value_d4)
        + int(value_e4)
        + int(value_f4)
        + int(value_g4)
        + int(value_h4)
        + int(value_i4)
        + int(value_j4)
        + int(value_k4)
        + int(value_m4)
        + int(value_n4)
        + int(value_o4)
        + int(value_p4)
    )

    total_names = ["kirmizi", "sari", "yesil", "mavi"]
    total = total_1 + total_2 + total_3 + total_4
    total1, total2, total3, total4 = (
        total_1 / total,
        total_2 / total,
        total_3 / total,
        total_4 / total,
    )
    rational = [total1, total2, total3, total4]
    color_distribution = zip(total_names, rational)
    color_distribution = sorted(color_distribution, key=lambda x: x[1], reverse=True)
    if round(color_distribution[0][1] * 10) / 10 >= 0.7:
        nihai_sonuc = nihai_sonuc_dict[color_distribution[0][0]]
        nihai_sonuc_aciklama = nihai_sonuc_aciklama_dict[color_distribution[0][0]]

    elif (
        color_distribution[0][0] == "kirmizi" and color_distribution[1][0] == "sari"
    ) or (color_distribution[0][0] == "sari" and color_distribution[1][0] == "kirmizi"):
        nihai_sonuc = nihai_sonuc_dict["kirmizi_sari"]
        nihai_sonuc_aciklama = nihai_sonuc_aciklama_dict["kirmizi_sari"]

    elif (
        color_distribution[0][0] == "mavi" and color_distribution[1][0] == "kirmizi"
    ) or (color_distribution[0][0] == "kirmizi" and color_distribution[1][0] == "mavi"):
        nihai_sonuc = nihai_sonuc_dict["mavi_kirmizi"]
        nihai_sonuc_aciklama = nihai_sonuc_aciklama_dict["mavi_kirmizi"]

    elif (
        color_distribution[0][0] == "yesil" and color_distribution[1][0] == "sari"
    ) or (color_distribution[0][0] == "sari" and color_distribution[1][0] == "yesil"):
        nihai_sonuc = nihai_sonuc_dict["yesil_sari"]
        nihai_sonuc_aciklama = nihai_sonuc_aciklama_dict["yesil_sari"]

    elif (
        color_distribution[0][0] == "mavi" and color_distribution[1][0] == "yesil"
    ) or (color_distribution[0][0] == "yesil" and color_distribution[1][0] == "mavi"):
        nihai_sonuc = nihai_sonuc_dict["mavi_yesil"]
        nihai_sonuc_aciklama = nihai_sonuc_aciklama_dict["mavi_yesil"]

    elif (
        color_distribution[0][0] == "kirmizi" and color_distribution[1][0] == "yesil"
    ) or (
        color_distribution[0][0] == "yesil" and color_distribution[1][0] == "kirmizi"
    ):
        nihai_sonuc = nihai_sonuc_dict["kirmizi_yesil"]
        nihai_sonuc_aciklama = nihai_sonuc_aciklama_dict["kirmizi_yesil"]

    elif (
        color_distribution[0][0] == "sari" and color_distribution[1][0] == "mavi"
    ) or (color_distribution[0][0] == "mavi" and color_distribution[1][0] == "sari"):
        nihai_sonuc = nihai_sonuc_dict["sari_mavi"]
        nihai_sonuc_aciklama = nihai_sonuc_aciklama_dict["sari_mavi"]

    else:
        nihai_sonuc = nihai_sonuc_dict[color_distribution[0][0]]
        nihai_sonuc_aciklama = nihai_sonuc_aciklama_dict[color_distribution[0][0]]


def versiyon():
    st.caption(
        """
                <p style='text-align: center;'>
                ver 1.2.0 <br/><font size="2">build 28112024.1628</font>
                </p>
            """,
        unsafe_allow_html=True,
    )
